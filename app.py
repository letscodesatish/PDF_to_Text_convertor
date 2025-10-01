# app.py
import streamlit as st
from pdf_to_images import pdf_to_images
from ocr_extraction import extract_text
from text_cleanup import clean_text
from docx import Document
from docx.shared import Pt
from io import BytesIO
from datetime import datetime

# -----------------------------
# Helper functions
# -----------------------------

def is_heading(line: str) -> bool:
    """
    Detects if a line is a heading.
    - Ends with ":" OR
    - Mostly uppercase
    """
    if line.endswith(":"):
        return True

    letters = [c for c in line if c.isalpha()]
    if not letters:
        return False

    uppercase_ratio = sum(1 for c in letters if c.isupper()) / len(letters)
    return uppercase_ratio >= 0.8

def process_pdf(pdf_file):
    """
    Converts PDF to images, extracts text using OCR, and cleans it.
    """
    images = pdf_to_images(pdf_file)
    text = ""
    for i, img in enumerate(images):
        page_text = extract_text(img)
        cleaned_text = clean_text(page_text)
        text += f"\n\n--- Page {i+1} ---\n{cleaned_text}"
    return text

def save_as_docx(text: str) -> BytesIO:
    """
    Converts text to a Word document in memory.
    Returns BytesIO object.
    """
    doc = Document()
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue

        p = doc.add_paragraph()
        run = p.add_run(line)

        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

        if is_heading(line):
            run.bold = True

    # Save to BytesIO
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# -----------------------------
# Streamlit UI
# -----------------------------
st.title("📄 PDF to Cleaned Word Document")
st.write("Upload your PDF and get a cleaned DOCX file with headings formatted.")

uploaded_file = st.file_uploader("Choose a PDF file", type=["pdf"])

if uploaded_file:
    with st.spinner("Processing PDF..."):
        text = process_pdf(uploaded_file)
        docx_file = save_as_docx(text)
    
    st.success("✅ Done! Your cleaned Word document is ready.")
    
    download_filename = f"cleaned_output_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    st.download_button(
        label="📥 Download DOCX",
        data=docx_file,
        file_name=download_filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
